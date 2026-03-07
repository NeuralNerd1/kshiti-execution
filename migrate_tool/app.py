from flask import Flask, render_template
from pathlib import Path
from flask import request, jsonify
import io
import zipfile
from pathlib import Path
from flask import send_file
from docx import Document
from docx.shared import Pt

from scanner import ProjectScanner
from utils.tree_builder import render_tree

app = Flask(__name__, template_folder="templates")

PROJECT_ROOT = Path(__file__).parent.parent.resolve()

def resolve_safe_path(relative_path: str) -> Path:
    """
    Prevent path traversal and ensure file stays within project root
    """
    resolved = (PROJECT_ROOT / relative_path).resolve()

    if not resolved.exists():
        raise FileNotFoundError(relative_path)

    if PROJECT_ROOT not in resolved.parents:
        raise PermissionError(f"Invalid path: {relative_path}")

    if resolved.is_dir():
        raise IsADirectoryError(relative_path)

    return resolved

def add_code_block(doc, code: str):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)

def add_heading(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    return h


def write_tree_to_doc(doc, tree: dict, base_path: Path, prefix=""):
    """
    Recursively write folders & files into Word document
    """
    for name, node in tree.items():
        current_path = f"{prefix}/{name}" if prefix else name

        # FILE
        if isinstance(node, dict) and {"path", "content", "skipped"} <= node.keys():
            doc.add_heading(f"📄 {current_path}", level=4)
            doc.add_paragraph(f"PATH: {current_path}")
            doc.add_paragraph("-" * 60)
            add_code_block(doc, node["content"])
            doc.add_paragraph("-" * 60)
            doc.add_page_break()
            continue

        # FOLDER
        doc.add_heading(f"📁 {current_path}", level=3)
        write_tree_to_doc(doc, node, base_path, current_path)



@app.route("/migrate/collect", methods=["POST"])
def collect_files():
    payload = request.get_json(force=True)
    files = payload.get("files", [])

    if not isinstance(files, list):
        return jsonify({"error": "Invalid payload"}), 400

    # TEMP: trust paths (Phase 5 will validate strictly)
    return jsonify({
        "count": len(files),
        "files": files
    })


@app.route("/migrate")
def migrate():
    root = Path(__file__).parent.parent  # project root
    scanner = ProjectScanner(root)
    tree = scanner.scan()

    tree_html = render_tree(tree)

    return render_template(
        "migrate.html",
        TREE_HTML=tree_html
    )

@app.route("/migrate/export/zip", methods=["POST"])
def export_zip():
    payload = request.get_json(force=True)
    files = payload.get("files", [])

    if not isinstance(files, list) or not files:
        return jsonify({"error": "No files provided"}), 400

    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        for rel_path in files:
            try:
                abs_path = resolve_safe_path(rel_path)
                zipf.write(
                    abs_path,
                    arcname=rel_path  # preserve hierarchy
                )
            except Exception as e:
                return jsonify({
                    "error": str(e),
                    "file": rel_path
                }), 400

    zip_buffer.seek(0)

    return send_file(
        zip_buffer,
        mimetype="application/zip",
        as_attachment=True,
        download_name="project_migration.zip"
    )

@app.route("/migrate/export/word", methods=["POST"])
def export_word():
    payload = request.get_json(force=True)
    files = payload.get("files", [])

    if not isinstance(files, list) or not files:
        return jsonify({"error": "No files provided"}), 400

    doc = Document()

    # Title
    title = doc.add_heading("PROJECT MIGRATION DOCUMENT", level=1)
    title.alignment = 1  # center

    doc.add_paragraph("")

    for rel_path in files:
        try:
            abs_path = resolve_safe_path(rel_path)
            content = abs_path.read_text(encoding="utf-8", errors="ignore")
        except Exception as e:
            return jsonify({"error": str(e), "file": rel_path}), 400

        # File header
        doc.add_heading(f"FILE: {rel_path}", level=3)
        doc.add_paragraph(f"PATH: {rel_path}")

        doc.add_paragraph("-" * 50)

        add_code_block(doc, content)

        doc.add_paragraph("-" * 50)
        doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="project_migration.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/migrate/export/word/full", methods=["GET"])
def export_full_project_word():
    scanner = ProjectScanner(PROJECT_ROOT)
    tree = scanner.scan()

    doc = Document()

    # Title
    title = doc.add_heading("PROJECT STRUCTURE EXPORT", level=1)
    title.alignment = 1  # center

    doc.add_paragraph(f"Root Directory: {PROJECT_ROOT.name}")
    doc.add_paragraph("=" * 70)

    write_tree_to_doc(doc, tree, PROJECT_ROOT)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="full_project_structure.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )




if __name__ == "__main__":
    app.run(
        port=5050,
        debug=True,
        use_reloader=False,
        threaded=False
    )
